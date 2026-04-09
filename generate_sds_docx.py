"""
Script to convert SDS.md to a properly formatted Word document
with a professional cover page for Final Year Project
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def create_cover_page(doc):
    """Create a professional cover page for Final Year Project SDS"""
    
    # Add cover page section
    cover_section = doc.add_section()
    
    # Title - Kingstone WiFi Billing System
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('KINGSTONE WIFI BILLING SYSTEM')
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.underline = True
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Software Design Specification (SDS)')
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.bold = True
    
    # Space
    doc.add_paragraph()
    
    # Project description
    proj_desc = doc.add_paragraph()
    proj_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    proj_desc_run = proj_desc.add_run('A Comprehensive WiFi Billing and Management System\nfor ISP Hotspots in Kenya')
    proj_desc_run.font.name = 'Arial'
    proj_desc_run.font.size = Pt(12)
    proj_desc_run.font.italic = True
    
    # Space
    for _ in range(3):
        doc.add_paragraph()
    
    # Student details
    details = [
        ('Student Name:', 'TRACY [YOUR FULL NAME]'),
        ('Registration Number:', '[YOUR REG NUMBER]'),
        ('Supervisor:', '[SUPERVISOR NAME]'),
        ('Department:', 'Computer Science/Information Technology'),
        ('Institution:', '[YOUR UNIVERSITY NAME]'),
    ]
    
    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = p.add_run(f'{label} ')
        label_run.font.name = 'Arial'
        label_run.font.size = Pt(12)
        label_run.font.bold = True
        value_run = p.add_run(value)
        value_run.font.name = 'Arial'
        value_run.font.size = Pt(12)
    
    # Space
    for _ in range(2):
        doc.add_paragraph()
    
    # Submission details
    submission = doc.add_paragraph()
    submission.alignment = WD_ALIGN_PARAGRAPH.CENTER
    submission_run = submission.add_run('Submitted in Partial Fulfillment of the Requirements\nfor the Degree of [DEGREE NAME]')
    submission_run.font.name = 'Arial'
    submission_run.font.size = Pt(12)
    
    # Space
    doc.add_paragraph()
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run('March 2026')
    date_run.font.name = 'Arial'
    date_run.font.size = Pt(14)
    date_run.font.bold = True
    
    # Add page break after cover
    doc.add_page_break()
    
    return doc


def create_title_page(doc):
    """Create the SDS title page"""
    
    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('SOFTWARE DESIGN SPECIFICATION')
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    
    # Project name
    proj_name = doc.add_paragraph()
    proj_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    proj_run = proj_name.add_run('Kingstone WiFi Billing System')
    proj_run.font.name = 'Arial'
    proj_run.font.size = Pt(16)
    proj_run.font.bold = True
    
    # Space
    doc.add_paragraph()
    
    # Document info table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    info = [
        ('Document Version:', '1.0.0'),
        ('Date:', 'March 31, 2026'),
        ('Prepared By:', 'Development Team'),
        ('Status:', 'Production Ready'),
        ('Document Type:', 'Software Design Specification'),
    ]
    
    for i, (label, value) in enumerate(info):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        
        p_label = cell_label.paragraphs[0]
        p_label.add_run(label)
        p_label.runs[0].font.bold = True
        p_label.runs[0].font.name = 'Arial'
        p_label.runs[0].font.size = Pt(11)
        
        p_value = cell_value.paragraphs[0]
        p_value.add_run(value)
        p_value.runs[0].font.name = 'Arial'
        p_value.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Copyright notice
    copyright_para = doc.add_paragraph()
    copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_run = copyright_para.add_run('© 2026 Kingstone WiFi Billing System\nAll Rights Reserved')
    copyright_run.font.name = 'Arial'
    copyright_run.font.size = Pt(10)
    copyright_run.font.italic = True
    
    # Add page break
    doc.add_page_break()
    
    return doc


def create_table_of_contents(doc):
    """Create table of contents"""
    
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_title.add_run('TABLE OF CONTENTS')
    toc_run.font.name = 'Arial'
    toc_run.font.size = Pt(16)
    toc_run.font.bold = True
    
    doc.add_paragraph()
    
    toc_items = [
        ('1. INTRODUCTION', '1'),
        ('1.1 Purpose', '1'),
        ('1.2 Scope', '1'),
        ('1.3 Definitions and Acronyms', '2'),
        ('1.4 References', '3'),
        ('1.5 Document Overview', '3'),
        ('2. SYSTEM OVERVIEW', '4'),
        ('2.1 System Context', '4'),
        ('2.2 System Functions', '5'),
        ('2.3 User Classes and Characteristics', '6'),
        ('2.4 Operating Environment', '7'),
        ('2.5 Design Constraints', '8'),
        ('2.6 Assumptions and Dependencies', '8'),
        ('3. ARCHITECTURAL DESIGN', '9'),
        ('3.1 High-Level Architecture', '9'),
        ('3.2 Technology Stack', '11'),
        ('3.3 Architecture Patterns', '12'),
        ('3.4 System Flow Diagrams', '13'),
        ('4. DATABASE DESIGN', '16'),
        ('4.1 Database Schema Overview', '16'),
        ('4.2 Table Specifications', '18'),
        ('4.3 Database Functions (RPC)', '25'),
        ('4.4 Data Integrity Constraints', '27'),
        ('5. INTERFACE DESIGN', '28'),
        ('5.1 User Interfaces', '28'),
        ('5.2 API Interfaces', '30'),
        ('5.3 Hardware Interfaces', '32'),
        ('5.4 Communication Interfaces', '33'),
        ('6. COMPONENT DESIGN', '34'),
        ('6.1 Frontend Components', '34'),
        ('6.2 Backend Components (Edge Functions)', '36'),
        ('6.3 Utility Components', '38'),
        ('7. SECURITY DESIGN', '39'),
        ('7.1 Authentication Security', '39'),
        ('7.2 Authorization Security', '41'),
        ('7.3 Data Protection', '42'),
        ('7.4 API Security', '43'),
        ('7.5 Audit Logging', '45'),
        ('8. INTEGRATION DESIGN', '46'),
        ('8.1 M-Pesa Integration', '46'),
        ('8.2 Mikrotik Integration', '48'),
        ('8.3 SMS Gateway Integration', '49'),
        ('8.4 Integration Error Handling', '50'),
        ('9. DEPLOYMENT DESIGN', '51'),
        ('9.1 Deployment Architecture', '51'),
        ('9.2 Environment Configuration', '52'),
        ('9.3 Deployment Steps', '53'),
        ('9.4 CI/CD Pipeline', '54'),
        ('9.5 Monitoring and Logging', '55'),
        ('10. TESTING STRATEGY', '56'),
        ('10.1 Testing Levels', '56'),
        ('10.2 Test Data Management', '57'),
        ('10.3 Performance Testing', '58'),
        ('10.4 Security Testing', '58'),
        ('10.5 Test Coverage Goals', '59'),
        ('11. APPENDICES', '60'),
        ('Appendix A: Database Migration Scripts', '60'),
        ('Appendix B: Edge Function Code', '60'),
        ('Appendix C: API Reference', '60'),
        ('Appendix D: Troubleshooting Guide', '61'),
        ('Appendix E: Glossary', '61'),
        ('Appendix F: Revision History', '62'),
        ('DOCUMENT APPROVAL', '63'),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item}')
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        
        # Add dots leader
        dots = '.' * (70 - len(item))
        p.add_run(f' {dots} {page}')
    
    doc.add_page_break()
    
    return doc


def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Arial'
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)
    return heading


def add_code_block(doc, code, language=None):
    """Add a formatted code block with monospace font"""
    # Create a paragraph with shading for code block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Add code with monospace font
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    # Add light gray shading (simulated by making it a distinct block)
    return p


def add_table_from_markdown(doc, markdown_table):
    """Parse and add a markdown table"""
    lines = markdown_table.strip().split('\n')
    if len(lines) < 3:
        return
    
    # Count columns
    header_parts = lines[0].split('|')
    num_cols = len([p for p in header_parts if p.strip()])
    
    # Create table
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = 'Table Grid'
    
    # Skip header separator line (|---|---|)
    data_lines = [l for l in lines if l.strip() and not re.match(r'^[\s|:-]+$', l)]
    
    for line in data_lines:
        row = table.add_row()
        cells = [c.strip() for c in line.split('|')]
        cells = [c for c in cells if c]  # Remove empty cells from edges
        
        for i, cell_text in enumerate(cells):
            if i < num_cols:
                cell = row.cells[i]
                p = cell.paragraphs[0]
                p.clear()
                p.add_run(cell_text)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
    
    return table


def parse_and_add_content(doc, content):
    """Parse markdown content and add to document"""
    lines = content.split('\n')
    i = 0
    current_heading_level = 0
    in_code_block = False
    code_block_content = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_block_content)
                if code_text.strip():
                    add_code_block(doc, code_text)
                code_block_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # Handle headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            
            # Skip main title (already in cover/title page)
            if level == 1 and 'SOFTWARE DESIGN SPECIFICATION' in text.upper():
                i += 1
                continue
            
            add_heading(doc, text, level=min(level, 4))
            i += 1
            continue
        
        # Handle horizontal rules (section dividers)
        if re.match(r'^---+$', line.strip()):
            i += 1
            continue
        
        # Handle tables
        if '|' in line and not line.startswith('|-'):
            # Check if this might be a table
            potential_table = []
            j = i
            while j < len(lines) and '|' in lines[j]:
                potential_table.append(lines[j])
                j += 1
            
            # If we have enough lines for a table
            if len(potential_table) >= 2:
                table_text = '\n'.join(potential_table)
                add_table_from_markdown(doc, table_text)
                i = j
                continue
        
        # Handle regular paragraphs
        if line.strip():
            # Check for bold/italic markdown
            p = doc.add_paragraph()
            
            # Process line for markdown formatting
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                elif part.strip():
                    run = p.add_run(part)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
        
        i += 1
    
    return doc


def main():
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Create cover page
    print("Creating cover page...")
    create_cover_page(doc)
    
    # Create title page
    print("Creating title page...")
    create_title_page(doc)
    
    # Create table of contents
    print("Creating table of contents...")
    create_table_of_contents(doc)
    
    # Read and parse SDS content
    print("Reading SDS.md...")
    with open('SDS.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse and add content
    print("Parsing and adding content...")
    parse_and_add_content(doc, content)
    
    # Add document approval page
    print("Adding approval page...")
    doc.add_page_break()
    add_heading(doc, 'DOCUMENT APPROVAL', level=1)
    
    approval_table = doc.add_table(rows=5, cols=4)
    approval_table.style = 'Table Grid'
    
    headers = ['Role', 'Name', 'Signature', 'Date']
    for i, header in enumerate(headers):
        cell = approval_table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        run.font.name = 'Arial'
    
    roles = [
        'Project Manager',
        'Technical Lead',
        'Quality Assurance',
        'System Administrator'
    ]
    
    for i, role in enumerate(roles):
        cell = approval_table.cell(i + 1, 0)
        p = cell.paragraphs[0]
        p.add_run(role)
        p.runs[0].font.name = 'Arial'
    
    # Save document
    print("Saving document...")
    doc.save('Kingstone_WiFi_Billing_SDS.docx')
    print("Done! Document saved as 'Kingstone_WiFi_Billing_SDS.docx'")


if __name__ == '__main__':
    main()
